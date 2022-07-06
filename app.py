from flask import Flask, request, redirect, render_template, url_for, flash
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from numpy import array
from docx import Document
from flask import send_from_directory
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_required, login_user, current_user, logout_user
from wtforms import StringField, SubmitField, TextAreaField, BooleanField, PasswordField
from wtforms.validators import DataRequired, Email
from flask_wtf import FlaskForm
import xlsxwriter
import os
import shutil
from os import path

app = Flask(__name__, static_folder='static')
app.debug = True
app.config['SECRET_KEY'] = 'a really really really really long secret key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///med.bd'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

consts = array(
    [-0.057, -0.043, 0.164, -0.103, -0.136, 0.176, -0.13, -0.039, 0.131, -0.076, -0.095, -0.256, 0.147, 0.029, 0.069,
     0.199, -0.303, 0.049, 0.046, -0.035, 0.105, -0.061, 0.503, -0.043, 0.268, -0.398, -0.047, -0.17, 0.412, -0.495,
     0.307, -0.241, 0.311, 0.315, 0.33], dtype='float')


@login_manager.user_loader
def load_user(user_id):
    return db.session.query(User).get(user_id)


class ContactForm(FlaskForm):
    name = StringField("Name: ", validators=[DataRequired()])
    message = TextAreaField("Message", validators=[DataRequired()])
    submit = SubmitField("Вход")


class LoginForm(FlaskForm):
    username = StringField("Логин", validators=[DataRequired()])
    password = PasswordField("Пароль", validators=[DataRequired()])
    remember = BooleanField("Запомнить меня")
    submit = SubmitField("Вход")


class User(db.Model, UserMixin):
    id = db.Column(db.Integer(), primary_key=True)
    name = db.Column(db.String(100))
    username = db.Column(db.String, nullable=False, unique=True)
    email = db.Column(db.String, nullable=False, unique=True)
    role = db.Column(db.String(50), nullable=False, default='user')
    password_hash = db.Column(db.String(100), nullable=False)
    created_on = db.Column(db.DateTime(), default=datetime.utcnow)
    updated_on = db.Column(db.DateTime(), default=datetime.utcnow, onupdate=datetime.utcnow)

    def __repr__(self):
        return "<{}:{}>".format(self.id, self.username)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)


class Patient(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.DateTime, default=datetime.utcnow)
    user = db.Column(db.String(100), nullable=False)
    med = db.Column(db.String(300), nullable=False)
    q1 = db.Column(db.String(100), nullable=False)
    q2 = db.Column(db.String(100), nullable=False)
    q3 = db.Column(db.Integer, nullable=False)
    q4 = db.Column(db.Integer, nullable=False)
    q5 = db.Column(db.Integer, nullable=False)
    q6 = db.Column(db.Integer, nullable=False)
    q7 = db.Column(db.Integer, nullable=False)
    q8 = db.Column(db.Integer, nullable=False)
    q9 = db.Column(db.Integer, nullable=False)
    q10 = db.Column(db.Integer, nullable=False)
    q11 = db.Column(db.Integer, nullable=False)
    q12 = db.Column(db.Integer, nullable=False)
    q13 = db.Column(db.Integer, nullable=False)
    q14 = db.Column(db.Integer, nullable=False)
    q15 = db.Column(db.Integer, nullable=False)
    q16 = db.Column(db.Integer, nullable=False)
    q17 = db.Column(db.Integer, nullable=False)
    q18 = db.Column(db.Integer, nullable=False)
    q19 = db.Column(db.Integer, nullable=False)
    q20 = db.Column(db.Integer, nullable=False)
    q21 = db.Column(db.Integer, nullable=False)
    q22 = db.Column(db.Integer, nullable=False)
    q23 = db.Column(db.Integer, nullable=False)
    q24 = db.Column(db.Integer, nullable=False)
    q25 = db.Column(db.Integer, nullable=False)
    q26 = db.Column(db.Integer, nullable=False)
    q27 = db.Column(db.Integer, nullable=False)
    q28 = db.Column(db.Integer, nullable=False)
    q29 = db.Column(db.Integer, nullable=False)
    q30 = db.Column(db.Integer, nullable=False)
    q31 = db.Column(db.Integer, nullable=False)
    q32 = db.Column(db.Integer, nullable=False)
    q33 = db.Column(db.Integer, nullable=False)
    q34 = db.Column(db.Integer, nullable=False)
    q35 = db.Column(db.Integer, nullable=False)
    q36 = db.Column(db.Integer, nullable=False)
    result = db.Column(db.Integer, nullable=False, default=0)


@app.route('/admin/')
@login_required
def admin():
    return render_template('admin.html')


@app.route('/login/', methods=['POST', 'GET'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        form = LoginForm()
        if form.validate_on_submit():
            user = db.session.query(User).filter(User.username == form.username.data).first()
            if user and user.check_password(form.password.data):
                login_user(user, remember=form.remember.data)
                return redirect(url_for('index'))

        flash("Неверный логин или пароль", 'error')
        return redirect(url_for('login'))
    return render_template('login.html', form=form)


@app.route('/logout/')
@login_required
def logout():
    logout_user()
    flash("Вы вышли из аккаунта")
    return redirect(url_for('login'))


@app.route('/history/')
@login_required
def history():
    patient = Patient.query.order_by(Patient.date.desc()).all()
    return render_template("history.html", patient=patient)


@app.route('/', methods=['POST','GET'])
@app.route('/index/', methods=['POST','GET'])
@login_required
def index():
    if request.method == "POST":
        name = request.form['x']
        patient = Patient.query.order_by(Patient.date.desc()).all()
        for el in patient:
            if name == el.q1:
                patientR = Patient.query.get(el.id)
                return render_template('result.html', patient=patientR, all_patient=patient)
            else:
                return render_template("index.html", otvet='Пациент не найден')
    return render_template("index.html")


@app.route('/quiz/', methods=['POST', 'GET'])
@login_required
def quiz():
    if request.method == 'POST':
        med = request.form['med']
        q1 = request.form['q1']
        q2 = request.form['q2']
        q3 = request.form['q3']
        q4 = request.form['q4']
        q5 = request.form['q5']
        q6 = request.form['q6']
        q7 = request.form['q7']
        q8 = request.form['q8']
        q9 = request.form['q9']
        q10 = request.form['q10']
        q11 = request.form['q11']
        q12 = request.form['q12']
        q13 = request.form['q13']
        q14 = request.form['q14']
        q15 = request.form['q15']
        q16 = request.form['q16']
        q17 = request.form['q17']
        q18 = request.form['q18']
        q19 = request.form['q19']
        q20 = request.form['q20']
        q21 = request.form['q21']
        q22 = request.form['q22']
        q23 = request.form['q23']
        q24 = request.form['q24']
        q25 = request.form['q25']
        q26 = request.form['q26']
        q27 = request.form['q27']
        q28 = request.form['q28']
        q29 = request.form['q29']
        q30 = request.form['q30']
        q31 = request.form['q31']
        q32 = request.form['q32']
        q33 = request.form['q33']
        q34 = request.form['q34']
        q35 = request.form['q35']
        q36 = request.form['q36']
        user = current_user.username
        params = array(
            [q3, q4, q5, q6, q7, q8, q9, q10, q11, q12, q13, q14, q15, q16, q17, q18, q19, q20, q21, q22, q23, q24, q25,
             q26, q27, q28, q29, q30, q31, q32,
             q33, q34, q35, q36], dtype='float')
        probability = min(1, max(0, consts[-1] + params.dot(consts[:-1])))

        patient = Patient(q1=q1, q2=q2, q3=q3, q4=q4, q5=q5, q6=q6, q7=q7, q8=q8, q9=q9, q10=q10, q11=q11, q12=q12,
                          q13=q13, q14=q14, q15=q15, q16=q16,
                          q17=q17, q18=q18, q19=q19, q20=q20, q21=q21, q22=q22, q23=q23, q24=q24, q25=q25, q26=q26,
                          q27=q27, q28=q28, q29=q29, q30=q30,
                          q31=q31, q32=q32, q33=q33, q34=q34, q35=q35, q36=q36, result=round(probability, 3), user=user, med=med)
        try:
            db.session.add(patient)
            db.session.commit()
            document = Document()
            table = document.add_table(rows=2, cols=4, style='Table Grid')
            hdr1_cells = table.rows[0].cells
            hdr1_cells[0].text = 'Дата обследования'
            hdr1_cells[1].text = 'Мед. учреждение'
            hdr1_cells[2].text = 'Пациент'
            hdr1_cells[3].text = 'Вероятность ИСМП'
            hdr2_cells = table.rows[1].cells
            hdr2_cells[0].text = str(patient.date)
            hdr2_cells[1].text = patient.med
            hdr2_cells[2].text = patient.q1
            hdr2_cells[3].text = "{:.3f}".format(patient.result)
            document.save('static/'+ str(current_user.username)+'/' + str(patient.q1) + '.docx')

            workbook = xlsxwriter.Workbook('static/'+ str(current_user.username)+'/' + str(patient.q1) + '.xlsx')
            bold = workbook.add_format({'bold': True})
            worksheet = workbook.add_worksheet()
            worksheet.write(0, 0, 'Дата обследования', bold)
            worksheet.write(0, 1, 'Мед. учреждение', bold)
            worksheet.write(0, 2, 'Пациент', bold)
            worksheet.write(0, 3, 'Вероятность ИСМП', bold)
            worksheet.write(1, 0, str(patient.date))
            worksheet.write(1, 1, patient.med)
            worksheet.write(1, 2, patient.q1)
            worksheet.write(1, 3, patient.result)
            workbook.close()
            all_patient = Patient.query.order_by(Patient.date.desc()).all()
            return render_template('result.html', patient=patient, all_patient=all_patient)
        except:
            return "При добавление данных произошла ошибка"
    else:
        return render_template('quiz.html')


@app.route('/all/')
@login_required
def vse():
    patient = Patient.query.order_by(Patient.date.desc()).all()
    document1 = Document()
    table = document1.add_table(rows=1, cols=4, style='Table Grid')
    hdr1_cells = table.rows[0].cells
    hdr1_cells[0].text = 'Дата обследования'
    hdr1_cells[1].text = 'Мед. учреждение'
    hdr1_cells[2].text = 'Пациент'
    hdr1_cells[3].text = 'Вероятность ИСМП'
    for el in patient:
        row_cells = table.add_row().cells
        row_cells[0].text = str(el.date)
        row_cells[1].text = el.med
        row_cells[2].text = el.q1
        row_cells[3].text = "{:.3f}".format(el.result)
    document1.save('static/all/all.docx')

    workbook = xlsxwriter.Workbook('static/all/all.xlsx')
    bold = workbook.add_format({'bold': True})
    worksheet = workbook.add_worksheet()
    worksheet.write(0, 0, 'Дата обследования', bold)
    worksheet.write(0, 1, 'Мед. учреждение', bold)
    worksheet.write(0, 2, 'Пациент', bold)
    worksheet.write(0, 3, 'Вероятность ИСМП', bold)
    i = 1
    for el in patient:
        worksheet.write(i, 0, str(el.date))
        worksheet.write(i, 1, el.med)
        worksheet.write(i, 2, el.q1)
        worksheet.write(i, 3, el.result)
        i = i+1
    workbook.close()
    return render_template("all.html", patient=patient)


@app.route('/all/<int:id_patient>', methods=['POST', 'GET'])
@login_required
def result(id_patient):
    patient = Patient.query.get(id_patient)
    all_patient = Patient.query.order_by(Patient.date.desc()).all()
    return render_template('result.html', patient=patient, all_patient=all_patient)


@app.route('/static/<path:filename>', methods=['GET', 'POST'])
@login_required
def download(filename):
    return send_from_directory(directory='static', filename=filename)


@app.route('/all_users/')
@login_required
def all_users():
    user = User.query.order_by().all()
    return render_template("all_users.html", user=user)


@app.route('/delete/<int:id>', methods=['GET', 'POST'])
@login_required
def delete(id):
    x = User.query.filter_by(id=id)
    y = User.query.get(id)
    shutil.rmtree("static/" + str(y.username))
    for i in x:
        db.session.delete(i)
    db.session.commit()
    user = User.query.order_by().all()
    return render_template('all_users.html', user=user)


@app.route('/create_user/', methods=['GET', 'POST'])
@login_required
def create_user():
    if request.method == 'POST':
        username = request.form['username']
        name = request.form['name']
        role = request.form['role']
        email = request.form['email']
        password_hash = generate_password_hash(request.form['password'])
        user = User(username=username, name=name, role=role, email=email, password_hash=password_hash)
        os.mkdir("static/" + str(user.username))
        try:
            db.session.add(user)
            db.session.commit()
            user = User.query.order_by().all()
            return render_template('all_users.html', user=user)
        except:
            return "При добавление данных произошла ошибка"

    return render_template("create_user.html")


if __name__ == '__main__':
    app.run(debug=True)
